"""Chuyển văn bản luật đã OCR thành .docx có phân cấp tiêu đề.

Nguồn: data/luat-116-2025-an-ninh-mang.pdf
    Luật số 116/2025/QH15 — LUẬT AN NINH MẠNG
    Tải từ cổng thông tin Chính phủ:
    https://datafiles.chinhphu.vn/cpp/files/vbpq/2026/01/luat116-2025.pdf

Đây là VĂN BẢN PHÁP LUẬT THẬT. Hai điều học viên phải biết:

1. Bản PDF gốc là bản SCAN (37 trang, 0 ký tự trích xuất được). Toàn bộ nội dung
   trong .docx đến từ OCR, nên CÓ SAI SÓT DẤU — ví dụ "kiêm tra" thay vì
   "kiểm tra". Đây là dữ liệu huấn luyện, KHÔNG dùng để tra cứu pháp lý.

2. Bản .docx được commit sẵn. Học viên không phải chạy OCR (bước đó chỉ chạy
   được trên macOS, xem scripts/ocr_pdf.py).

    uv run python scripts/build_corpus.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))
from rag_qa import config  # noqa: E402

NGUON_TXT = ROOT / "data" / "luat-116-2025-an-ninh-mang.txt"

CANH_BAO = (
    "VĂN BẢN DÙNG CHO ĐÀO TẠO KỸ THUẬT. Nội dung được OCR từ bản scan nên có "
    "sai sót về dấu và ký tự. KHÔNG dùng bản này để tra cứu hoặc trích dẫn "
    "pháp lý — hãy dùng bản gốc tại cổng thông tin Chính phủ."
)


def _la_so_trang(dong: str) -> bool:
    """Số trang do OCR đọc được, cần bỏ đi."""
    return bool(re.fullmatch(r"\d{1,3}", dong.strip()))


def build() -> Path:
    if not NGUON_TXT.exists():
        raise SystemExit(
            f"Chưa có {NGUON_TXT}.\n"
            "Chạy OCR trước (chỉ trên macOS):\n"
            "    uv run --with pypdfium2 --with pyobjc-framework-Vision "
            "--with pyobjc-framework-Quartz python scripts/ocr_pdf.py"
        )

    doc = Document()
    doc.add_heading("Luật số 116/2025/QH15 — Luật An ninh mạng", level=0)
    doc.add_paragraph(CANH_BAO)
    doc.add_paragraph(
        "Nguồn: https://datafiles.chinhphu.vn/cpp/files/vbpq/2026/01/luat116-2025.pdf"
    )

    so_chuong = so_dieu = 0
    for dong in NGUON_TXT.read_text(encoding="utf-8").split("\n"):
        dong = dong.strip()
        if not dong or _la_so_trang(dong):
            continue

        if re.match(r"^Chương [IVXLC]+", dong):
            doc.add_heading(dong, level=1)
            so_chuong += 1
        elif re.match(r"^Điều \d+\.", dong):
            doc.add_heading(dong, level=2)
            so_dieu += 1
        else:
            doc.add_paragraph(dong)

    config.DATA_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(config.CORPUS_PATH))
    print(f"Đã tạo: {config.CORPUS_PATH}")
    print(f"  {so_chuong} chương, {so_dieu} điều")
    return config.CORPUS_PATH


if __name__ == "__main__":
    build()
